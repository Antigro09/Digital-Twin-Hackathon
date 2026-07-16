from __future__ import annotations

import base64
import binascii
import csv
import io
import json
import zipfile
from dataclasses import dataclass

from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader

from .errors import DomainError
from .intelligence_models import DocumentImportRequest
from .safety import content_sha256


TEXT_MEDIA_TYPES = frozenset(
    {
        "text/plain",
        "text/markdown",
        "text/csv",
        "application/csv",
        "application/json",
        "application/yaml",
        "text/yaml",
        "application/xml",
        "text/xml",
        "image/svg+xml",
        "text/vnd.mermaid",
    }
)
PDF = "application/pdf"
DOCX = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
XLSX = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
SUPPORTED_MEDIA_TYPES = TEXT_MEDIA_TYPES | {PDF, DOCX, XLSX}

_MAX_OFFICE_ARCHIVE_ENTRIES = 10_000
_MAX_OFFICE_UNCOMPRESSED_BYTES = 50 * 1024 * 1024
_MAX_OFFICE_ENTRY_BYTES = 20 * 1024 * 1024
_MAX_OFFICE_COMPRESSION_RATIO = 200
_MAX_EXTRACTED_TEXT_BYTES = 20 * 1024 * 1024


@dataclass(frozen=True)
class ParsedDocument:
    text: str
    parser: str
    content_sha256: str
    raw_bytes: int


def _decode_utf8(raw: bytes) -> str:
    try:
        return raw.decode("utf-8")
    except UnicodeDecodeError as exc:
        raise DomainError(
            "invalid_document_encoding",
            "Text documents must use UTF-8 encoding.",
            status_code=422,
        ) from exc


def _pdf(raw: bytes) -> str:
    try:
        reader = PdfReader(io.BytesIO(raw), strict=True)
        if reader.is_encrypted:
            raise DomainError(
                "encrypted_document_not_supported",
                "Encrypted PDF documents are not supported.",
                status_code=422,
            )
        if len(reader.pages) > 500:
            raise DomainError("document_limit_exceeded", "PDF exceeds 500 pages.", status_code=413)
        return "\n\n".join((page.extract_text() or "") for page in reader.pages)
    except DomainError:
        raise
    except Exception as exc:
        raise DomainError("invalid_document", "The PDF could not be parsed safely.", status_code=422) from exc


def _validate_office_archive(raw: bytes) -> None:
    """Reject encrypted or expansion-heavy OOXML before a parser opens it."""

    try:
        with zipfile.ZipFile(io.BytesIO(raw)) as archive:
            entries = archive.infolist()
            if not entries or len(entries) > _MAX_OFFICE_ARCHIVE_ENTRIES:
                raise DomainError(
                    "document_archive_limit_exceeded",
                    "The Office document archive has an unsafe entry count.",
                    status_code=413,
                )
            total_uncompressed = 0
            for entry in entries:
                if entry.flag_bits & 0x1:
                    raise DomainError(
                        "encrypted_document_not_supported",
                        "Encrypted Office documents are not supported.",
                        status_code=422,
                    )
                if entry.file_size > _MAX_OFFICE_ENTRY_BYTES:
                    raise DomainError(
                        "document_archive_limit_exceeded",
                        "The Office document archive contains an oversized entry.",
                        status_code=413,
                    )
                total_uncompressed += entry.file_size
                if total_uncompressed > _MAX_OFFICE_UNCOMPRESSED_BYTES:
                    raise DomainError(
                        "document_archive_limit_exceeded",
                        "The Office document archive expands beyond the supported limit.",
                        status_code=413,
                    )
                compressed = max(1, entry.compress_size)
                if (
                    entry.file_size > 1024 * 1024
                    and entry.file_size > compressed * _MAX_OFFICE_COMPRESSION_RATIO
                ):
                    raise DomainError(
                        "document_archive_limit_exceeded",
                        "The Office document archive has an unsafe compression ratio.",
                        status_code=413,
                    )
    except DomainError:
        raise
    except (zipfile.BadZipFile, OSError) as exc:
        raise DomainError(
            "invalid_document",
            "The Office document archive could not be parsed safely.",
            status_code=422,
        ) from exc


def _docx(raw: bytes) -> str:
    try:
        _validate_office_archive(raw)
        document = Document(io.BytesIO(raw))
        values = [paragraph.text for paragraph in document.paragraphs]
        for table in document.tables[:200]:
            for row in table.rows[:5000]:
                values.append("\t".join(cell.text for cell in row.cells))
        return "\n".join(values)
    except DomainError:
        raise
    except Exception as exc:
        raise DomainError("invalid_document", "The DOCX file could not be parsed safely.", status_code=422) from exc


def _xlsx(raw: bytes) -> str:
    try:
        _validate_office_archive(raw)
        workbook = load_workbook(io.BytesIO(raw), read_only=True, data_only=True)
        lines: list[str] = []
        cells = 0
        for worksheet in workbook.worksheets[:50]:
            lines.append(f"[sheet:{worksheet.title}]")
            for row in worksheet.iter_rows(values_only=True):
                cells += len(row)
                if cells > 100_000:
                    raise DomainError(
                        "document_limit_exceeded",
                        "Spreadsheet exceeds 100000 cells.",
                        status_code=413,
                    )
                lines.append("\t".join("" if value is None else str(value) for value in row))
        workbook.close()
        return "\n".join(lines)
    except DomainError:
        raise
    except Exception as exc:
        raise DomainError(
            "invalid_document", "The XLSX file could not be parsed safely.", status_code=422
        ) from exc


def parse_document(body: DocumentImportRequest, *, max_bytes: int) -> ParsedDocument:
    media_type = body.media_type.split(";", 1)[0].strip().casefold()
    if media_type not in SUPPORTED_MEDIA_TYPES:
        raise DomainError(
            "unsupported_document_type",
            "The document media type is not supported.",
            status_code=415,
        )
    if body.text is not None:
        if media_type not in TEXT_MEDIA_TYPES:
            raise DomainError(
                "binary_document_requires_base64",
                "Binary documents must be supplied as base64 for safe parsing.",
                status_code=422,
            )
        raw = body.text.encode("utf-8")
        text = body.text
        parser = "trusted-transport-utf8/1.0"
    else:
        try:
            raw = base64.b64decode(body.content_base64 or "", validate=True)
        except (binascii.Error, ValueError) as exc:
            raise DomainError(
                "invalid_document_encoding",
                "content_base64 is not valid base64.",
                status_code=422,
            ) from exc
        if len(raw) > max_bytes:
            raise DomainError("document_limit_exceeded", "Document exceeds the byte limit.", status_code=413)
        if media_type in TEXT_MEDIA_TYPES:
            text = _decode_utf8(raw)
            parser = "utf8/1.0"
        elif media_type == PDF:
            text = _pdf(raw)
            parser = "pypdf/6.10.2"
        elif media_type == DOCX:
            text = _docx(raw)
            parser = "python-docx/1.2.0"
        elif media_type == XLSX:
            text = _xlsx(raw)
            parser = "openpyxl/3.1.5"
        else:  # pragma: no cover - guarded by the closed set above.
            raise DomainError("unsupported_document_type", "Unsupported document.", status_code=415)
    if len(raw) > max_bytes:
        raise DomainError("document_limit_exceeded", "Document exceeds the byte limit.", status_code=413)
    if not text.strip():
        raise DomainError(
            "document_contains_no_text",
            "The document parser did not extract any text.",
            status_code=422,
        )
    if len(text.encode("utf-8")) > _MAX_EXTRACTED_TEXT_BYTES:
        raise DomainError(
            "document_extraction_limit_exceeded",
            "The extracted document text exceeds the supported limit.",
            status_code=413,
        )
    # Validate syntactic JSON/CSV when the caller declares those formats. This
    # prevents mislabeled arbitrary bytes from entering the retrieval corpus.
    if media_type == "application/json":
        try:
            json.loads(text)
        except json.JSONDecodeError as exc:
            raise DomainError("invalid_document", "The JSON document is invalid.", status_code=422) from exc
    if media_type in {"text/csv", "application/csv"}:
        try:
            next(csv.reader(io.StringIO(text)), None)
        except csv.Error as exc:
            raise DomainError("invalid_document", "The CSV document is invalid.", status_code=422) from exc
    return ParsedDocument(
        text=text,
        parser=parser,
        content_sha256=content_sha256(raw),
        raw_bytes=len(raw),
    )
